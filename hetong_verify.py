import xlrd
import docx
import re
from win32com import client as wc
import os
from time import sleep

'''
    需求说明：
        需要检查100多份合同里面的数据是否填错，根据的数据是从数据库查询出来然后放在Excel表格中的。
        如果手动去对比需要一个一个查看很麻烦，所以写个简单的自动化脚本。
    实现步骤：
        1.把需要检查的数据从Excel表格中提取出来，放在list列表中。
        2.把合同的数据根据正则匹配提取出来，也放在list列表中。
        3.然后对比两个list是否相同
        4.如果不同就输出把不同的找出来。
'''

class OpenExcel():
    '''
    读取excel文件
    '''

    def __init__(self, path):
        # 打开excel文件
        self.data = xlrd.open_workbook(path)
        self.sheet = self.data.sheet_by_index(0)
    # 获取单元格的值
    def get_table_value(self, *args):
        # 获取读入的文件的

        a, b = args[0]
        data = self.sheet.cell(a, b).value
        return data


class OpenDocx():
    '''
    读取docx文件
    '''
    # 将doc转为docx:
    def doc_convert_docx(self, doc_path, docx_path):
        # path = r'E:\代偿用户对比\对比文件\委托代偿协议.doc'
        # path1 = r'E:\代偿用户对比\对比文件\委托代偿协议.docx'
        word = wc.Dispatch("Word.Application")
        # 打开被转换的文件
        doc = word.Documents.Open(doc_path)
        # 转换后要保存的路径及文件名
        doc.SaveAs(docx_path, 12) # 12为docx
        # 关闭文件
        doc.Close()
        # 退出
        word.Quit()

    # 获取文件内容
    def get_docx_file(self,docx_path):
        # 读取docx文件
        data = docx.Document(docx_path)
        t = []
        # 循环读取每一行内容，并把他加入列表中
        for i in range(0, 7):
            a = data.paragraphs[i]
            t.append(a.text)

        tables = data.tables  # 获取文件中的表格集
        table = tables[0]  # 获取文件中的第一个表格
        for i in range(1, len(table.rows)):  # 从表格第二行开始循环读取表格数据
            result = table.cell(i, 0).text + "" + table.cell(i, 1).text
            t.append(result)
        # 将列表的内容转换成字符串
        s = "".join(t)
        s = s.replace(' ', '')
        #print(s)
        # 正则表达式
        r = '合同编号：(.*)借款人（甲方）：(.*)（用户名：(.*?)）.*证件号码：(.*)联系电话：(\d+)借款本金数额（小写）￥(\d+).*人民币借款期限(.*)；.*借款利率(.*)/年.*'
        # 匹配内容
        p = re.search(r, s)
        user_data2 = list(p.groups())
        return user_data2


def file_name(file_dir):
    '''
    获取目录的所有文件名
    :param file_dir:
    :return:
    '''
    for root, dirs, files in os.walk(file_dir):

        # print('当前目录路径',root)  # 当前目录路径
        # print('当前路径下所有子目录',dirs)  # 当前路径下所有子目录
        # print('当前路径下所有非目录子文件',files)  # 当前路径下所有非目录子文件
        return dirs , files

def save_result(data):
    '''
    将对比结果写入文件
    :param data:
    :return:
    '''
    with open(r'D:\借款合同对比结果.txt', 'a') as f:
        f.write(data)

def main():
    # 单元格数据名称
    a = ['合同编号:','甲方:','用户名','身份证号:','手机号:', '借款金额:', '借款期限','利率']
    # 要读取的单元格数据
    cell = [(1,6),(1,3),(1,8),(1,4),(1,5),(1,1),(7,6),(1,9)]
    path = r'E:\代偿用户对比\对比文件\表格'
    path2 = r'E:\代偿用户对比\对比文件\合同'
    # 获取表格目录下的所有文件名
    tible_file = file_name(path)[1]
    # 获取合同目录下的所有文件名
    hetong_file = file_name(path2)[0]
    # 统计执行了几次
    count = 1
    # 统计不一样的有几次
    bt = 1

    # 对比表格目录和合同目录的文件文件名是否一样
    for table_username in tible_file:
        username1 = table_username.split('.')
        username1 = username1[0]
        for hetong_username in hetong_file:
            # 如果文件名相同就进行数据对比
            if username1 == hetong_username :
                excel_file = r'E:\代偿用户对比\对比文件\表格' + '\\' + table_username
                docx_file = r'E:\代偿用户对比\对比文件\合同' + '\\' + username1 + '\\' + '借款合同.docx'
                # print(excel_file)
                # print(doc_file)
                # print(docx_file)
                try:
                    user_data1 = []
                    # 用来判断是否是最后两个单元格数据 如果是需要做特殊处理
                    num = len(cell)-1
                    for data in cell:
                        # 提取cell列表相对应的单元格的值
                        d = OpenExcel(excel_file).get_table_value(data)
                        # 判断某个单元格的值是否是0.0如果是的话就转换成0
                        if d == 0.0 or d == '':
                            d = '0'
                        # 判断是否是最后两个数值，如果是的话就乘以100
                        if num == 0 and type(d)== float:
                            d = round(d * 100, 2)
                            d = str(d) + '%'

                        # 判断某个单元格的值是否是小数，如果是的话就保留两位小数四舍五入
                        if type(d) == float:
                            d = round(d, 2)
                            # 最后小数位是否为0如果是的话就把0去掉
                            if int(d) - d == 0:
                                d = int(d)
                            d = str(d)
                        # 最后吧获取到的时候保存到表格
                        user_data1.append(d)
                        num -= 1


                    # 修改doc为docx
                    # OpenDocx(doc_path=doc_file, docx_path=docx_file).doc_convert_docx()
                    user_data2 = OpenDocx().get_docx_file(docx_file)


                    # print(user_data1)
                    # print(user_data2)

                    num1 = 0
                    d1 = []
                    # 把单元格名称和单元格数据进行拼接
                    for i in a:
                        d1.append(i + user_data1[num1])
                        num1 += 1
                    print(d1)

                    num2 = 0
                    d2 = []
                    # 把名称和提取到的合同数据进行拼接
                    for i in a:
                        d2.append(i+user_data2[num2])
                        num2 += 1
                    print(d2)
                    # 输出两个列表不同的数据 差集
                    print(list(set(d1).difference(set(d2))))
                    print(list(set(d2).difference(set(d1))))
                    print(d1 == d2)
                    b = d1==d2
                    # 把不相等的数据写入到文件
                    if b == False and len(list(set(d1).difference(set(d2)))) > 0 :
                        bt_d1 = ",".join(d1)
                        bt_d2 = ",".join(d2)
                        bt_d1_result = ",".join(list(set(d1).difference(set(d2))))
                        bt_d2_result = ",".join(list(set(d2).difference(set(d1))))
                        save_result(bt_d1 + '\n')
                        save_result(bt_d2 + '\n\n')
                        save_result('不同的数据如下：' + '\n')
                        save_result('表格文件数据：\t'+ bt_d1_result + '\n\n')
                        save_result('借款合同数据：\t' + bt_d2_result + '\n\n')
                        save_result('========================================================' + '\n')
                        print('不相同的有%d个' % bt)
                        bt += 1
                    print('========================== %d ================================' % count)
                    count += 1
                    sleep(2)
                except Exception as msg:
                    print(msg)
                    print(username1, '没有对比成功')
main()